"""Document bytes -> plain text.

Deliberately narrow: this layer only turns a file into text and reports honestly when it
cannot. Interpreting that text is `cv_parser`'s job, and deciding what to do with the result
is yours.

A CV that is a scanned image extracts to nothing. That is reported as an explicit failure
rather than silently producing an empty profile -- "we found no facts" and "we could not
read the file" are very different messages to show someone.
"""

from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass

PDF_MAGIC = b"%PDF-"
ZIP_MAGIC = b"PK\x03\x04"  # docx is a zip

#: Below this, a "successful" extraction is almost certainly a scanned or empty document.
MIN_USEFUL_CHARS = 120


class ExtractionError(RuntimeError):
    """The file could not be turned into text. The message is shown to the user."""


@dataclass
class Extracted:
    text: str
    parser: str
    page_count: int | None = None

    @property
    def chars(self) -> int:
        return len(self.text)


def sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sniff(data: bytes, filename: str = "") -> str:
    """Identify the format from content first, extension second.

    Content wins because a mislabelled extension is common and a wrong parser produces
    confusing garbage rather than a clean error.
    """
    if data.startswith(PDF_MAGIC):
        return "pdf"
    if data.startswith(ZIP_MAGIC):
        return "docx"
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return "pdf"
    if lowered.endswith((".docx", ".doc")):
        return "docx"
    if lowered.endswith((".txt", ".md", ".rtf")):
        return "text"
    # Fall back to text only if it decodes cleanly as UTF-8.
    try:
        data.decode("utf-8")
    except UnicodeDecodeError:
        return "unknown"
    return "text"


def _normalise(text: str) -> str:
    """Tidy extractor output without destroying line structure.

    Line breaks matter: CV parsing leans on them to tell a heading from its content, so
    only runs of blank lines and trailing whitespace are collapsed.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ").replace("​", "")
    # Ligatures that PDF extraction commonly emits.
    for bad, good in (("ﬁ", "fi"), ("ﬂ", "fl"), ("ﬀ", "ff")):
        text = text.replace(bad, good)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _extract_pdf(data: bytes) -> Extracted:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ExtractionError("PDF support is not installed (pypdf).") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # An empty password unlocks many "protected" PDFs; if it does not, say so.
            try:
                reader.decrypt("")
            except Exception as exc:  # noqa: BLE001
                raise ExtractionError(
                    "This PDF is password-protected. Save an unprotected copy and upload that."
                ) from exc
        pages = [page.extract_text() or "" for page in reader.pages]
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 - pypdf raises a wide variety
        raise ExtractionError(f"Could not read this PDF ({exc.__class__.__name__}).") from exc

    return Extracted(text=_normalise("\n".join(pages)), parser="pdf", page_count=len(pages))


def _extract_docx(data: bytes) -> Extracted:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("DOCX support is not installed (python-docx).") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(
            "Could not read this file as a .docx. If it is an old .doc, re-save it as .docx."
        ) from exc

    parts = [p.text for p in document.paragraphs]
    # Tables carry real content in many CV templates -- skills grids, date columns.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  ".join(cells))

    return Extracted(text=_normalise("\n".join(parts)), parser="docx")


def _extract_text(data: bytes) -> Extracted:
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return Extracted(text=_normalise(data.decode(encoding)), parser="text")
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode this file as text.")


def extract(data: bytes, filename: str = "") -> Extracted:
    """Turn an uploaded file into text, or raise `ExtractionError` explaining why not."""
    if not data:
        raise ExtractionError("The file is empty.")

    kind = sniff(data, filename)
    if kind == "pdf":
        result = _extract_pdf(data)
    elif kind == "docx":
        result = _extract_docx(data)
    elif kind == "text":
        result = _extract_text(data)
    else:
        raise ExtractionError(
            f"Unsupported file type for {filename or 'this upload'}. "
            "Upload a PDF, a .docx, or plain text."
        )

    if result.chars < MIN_USEFUL_CHARS:
        raise ExtractionError(
            f"Only {result.chars} characters of text came out of this file. "
            "It is probably a scan or an image-only export -- upload a text-based version, "
            "since no OCR runs here."
        )
    return result
